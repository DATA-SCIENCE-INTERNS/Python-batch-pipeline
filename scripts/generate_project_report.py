"""Generate the formal NYC Taxi Batch Pipeline project report."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "NYC_Taxi_Batch_Pipeline_Project_Report.docx"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, True, WHITE, 9)
        shade(table.rows[0].cells[index], NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, False, None, 8.5)
            if row_index % 2:
                shade(cells[index], LIGHT_GREY)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "1E1E1E")
    cell.margin_top = Inches(0.08)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.strip().splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(230, 230, 230)
    doc.add_paragraph()


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select Update Field to generate the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, placeholder, end])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 30, NAVY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, NAVY),
    ]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
        styles[name].font.bold = True

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.style = styles["Normal"]
        footer.add_run("NYC Taxi Batch Pipeline | Project Report   ")
        add_page_number(footer)

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)


def build_report():
    doc = Document()
    configure_document(doc)

    # Cover page
    doc.add_paragraph("DATA ENGINEERING PROJECT", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph("NYC Taxi Batch Pipeline", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Design, Implementation, Testing, Operations, Results and Critical Evaluation"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph("\n\n")
    cover = doc.add_table(rows=5, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.style = "Table Grid"
    cover_data = [
        ("Project type", "Python batch data pipeline"),
        ("Data source", "NYC TLC Yellow and Green Taxi Trip Records"),
        ("Architecture", "Bronze / Silver / Gold (Medallion)"),
        ("Technology", "Python, PyArrow, Pandas, PostgreSQL, Docker Compose"),
        ("Report date", date.today().strftime("%d %B %Y")),
    ]
    for i, (key, value) in enumerate(cover_data):
        set_cell_text(cover.cell(i, 0), key, True, WHITE, 10)
        shade(cover.cell(i, 0), NAVY)
        set_cell_text(cover.cell(i, 1), value, False, None, 10)
    doc.add_paragraph("\n")
    p = doc.add_paragraph("Prepared by: ______________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Institution / Organisation: ______________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Document Control", level=1)
    add_table(doc, ["Item", "Details"], [
        ("Document", "NYC Taxi Batch Pipeline Project Report"),
        ("Version", "1.0"),
        ("Status", "Implementation and evaluation report"),
        ("Scope", "2025 Yellow and Green monthly taxi batches"),
        ("Repository", "Python-batch-pipeline"),
    ], [1.7, 4.9])
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This project implements a reproducible Python batch ingestion pipeline for monthly "
        "NYC Taxi and Limousine Commission (TLC) Yellow and Green trip data. The system "
        "downloads Parquet files, preserves raw source artifacts, normalizes two different "
        "source schemas, validates each record, stores accepted and rejected data separately, "
        "and publishes deduplicated analytical tables in PostgreSQL. Docker Compose provides "
        "a repeatable local execution environment."
    )
    doc.add_paragraph(
        "The solution demonstrates core data-engineering concerns: bounded-memory processing, "
        "data contracts, lineage, checksums, audit metadata, transactional consistency, "
        "backfilling, failure recovery and idempotency. A deterministic trip key and PostgreSQL "
        "primary-key enforcement ensure that reingesting the same source does not duplicate Gold "
        "records. Invalid rows are quarantined with explicit rejection reasons instead of being "
        "silently discarded."
    )
    add_callout(doc, "Verified outcome", "Green Taxi data was loaded for all twelve months of 2025. Yellow Taxi data was loaded for January through March 2025. A Green January rerun promoted zero additional Gold rows, demonstrating idempotency.", GREEN)
    doc.add_paragraph(
        "The project is suitable as a portfolio, demonstration or learning pipeline, but it is "
        "not yet production-scale. The main technical limitation is the use of one large "
        "transaction per monthly batch, which produces heavy PostgreSQL write-ahead-log (WAL) "
        "activity and requires the entire month to be retried after an interruption."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph(
        "NYC TLC publishes large monthly Parquet datasets for several taxi services. Yellow and "
        "Green Taxi files contain related trip information but do not use identical field names "
        "or optional columns. Loading these files consistently requires controlled extraction, "
        "schema normalization, validation, lineage and repeatable database publication."
    )
    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Raw monthly files alone are difficult to query reliably. They can have schema "
        "differences, invalid records, repeated ingestion attempts and partial availability. The "
        "project therefore needed a batch process that could preserve source data, produce a "
        "canonical dataset, reject bad records transparently, avoid duplicate publication and "
        "record evidence for every run."
    )
    doc.add_heading("1.3 Objectives", level=2)
    add_bullets(doc, [
        "Ingest monthly Yellow and Green Taxi Parquet files from the official NYC TLC source.",
        "Process large files in bounded-memory chunks rather than loading a whole month at once.",
        "Normalize source-specific fields into one canonical Silver schema.",
        "Apply explicit data-quality rules and preserve rejected records for investigation.",
        "Publish trusted, deduplicated Gold tables and a monthly reporting summary.",
        "Make reruns idempotent and maintain file- and run-level lineage.",
        "Support single-month ingestion and historical date-range backfills through a CLI.",
        "Provide repeatable local infrastructure, tests, documentation and operational queries.",
    ])
    doc.add_heading("1.4 Scope", level=2)
    add_table(doc, ["In scope", "Out of scope / future work"], [
        ("Yellow and Green monthly trip files", "Real-time or streaming ingestion"),
        ("Local Bronze filesystem", "Cloud object storage and lifecycle policies"),
        ("PostgreSQL Silver and Gold layers", "Distributed warehouse or lakehouse deployment"),
        ("CLI-driven monthly and range processing", "Managed scheduling and alerting"),
        ("Unit tests and observed integration evidence", "Full automated integration/chaos test suite"),
    ])

    doc.add_heading("2. Technology Stack", level=1)
    add_table(doc, ["Technology", "Role", "Reason for selection"], [
        ("Python 3.12+", "Pipeline control", "Readable ecosystem and strong data tooling"),
        ("Requests", "HTTP extraction", "Streaming downloads and explicit timeouts"),
        ("PyArrow", "Parquet access", "Metadata inspection and bounded-size record batches"),
        ("Pandas", "Transformation", "Schema mapping, explicit conversion and validation"),
        ("PostgreSQL 16", "Silver, Gold and metadata", "Transactions, constraints, SQL analytics and persistence"),
        ("psycopg2", "Database client", "Parameterized SQL and transaction control"),
        ("Click", "Command-line interface", "Structured ingest and backfill commands"),
        ("Docker Compose", "Local infrastructure", "Repeatable database and pipeline services"),
        ("Pytest", "Automated tests", "Fast unit testing of pipeline behavior"),
    ], [1.25, 1.65, 3.8])

    doc.add_heading("3. Solution Architecture", level=1)
    doc.add_heading("3.1 High-Level Data Flow", level=2)
    add_code(doc, """
NYC TLC monthly Parquet source
              |
              v
Bronze filesystem (raw, partitioned, checksummed)
              |
              v
PyArrow chunk reader -> Pandas canonical normalization
              |
              v
       Data-quality validation
          /                 \\
         v                   v
Silver validated trips   pipeline.rejected_records
         |
         v
Gold Yellow / Green tables (trip_key primary key)
         |
         v
gold.monthly_summary materialized view

Cross-cutting audit: pipeline.batch_runs + pipeline.file_ingestions
""")
    doc.add_heading("3.2 Medallion Layers", level=2)
    add_table(doc, ["Layer", "Storage", "Responsibility"], [
        ("Bronze", "Partitioned Parquet files", "Preserve source data and enable safe replay"),
        ("Silver", "silver.trips", "Canonical, typed and validated records with lineage"),
        ("Gold", "gold.yellow_trips / green_trips", "Trusted, deduplicated analytical records"),
        ("Reporting", "gold.monthly_summary", "Precomputed trip, distance and revenue summaries"),
    ])
    doc.add_heading("3.3 Deployment Topology", level=2)
    doc.add_paragraph(
        "Docker Compose runs PostgreSQL as a persistent service and the Python pipeline as an "
        "on-demand tool container. Containers communicate using the Compose hostname "
        "'postgres' on internal port 5432. The database is published to Windows on port 5433 "
        "because a native PostgreSQL service already occupies host port 5432."
    )

    doc.add_heading("4. Pipeline Design and Implementation", level=1)
    doc.add_heading("4.1 Extraction and Bronze Storage", level=2)
    add_numbered(doc, [
        "Construct the official monthly URL from taxi type, year and month.",
        "Create a partitioned Bronze path by taxi type, year and month.",
        "Reuse an existing file unless overwrite is requested.",
        "Stream the HTTP response to a temporary .part file.",
        "Reject HTTP errors and empty downloads; remove incomplete temporary files.",
        "Atomically rename the completed file and calculate its SHA-256 checksum.",
    ])
    add_code(doc, "data/bronze/{taxi_type}/{year}/{month}/{taxi_type}_tripdata_YYYY-MM.parquet")

    doc.add_heading("4.2 Schema Normalization", level=2)
    doc.add_paragraph(
        "Yellow and Green files use different timestamp names (tpep versus lpep) and contain "
        "different optional fields. The transformation module maps both sources to a shared "
        "canonical schema, inserts nullable columns where a source does not provide them and "
        "performs explicit numeric and datetime conversions."
    )
    add_table(doc, ["Source difference", "Canonical handling"], [
        ("tpep_pickup_datetime / lpep_pickup_datetime", "pickup_datetime"),
        ("tpep_dropoff_datetime / lpep_dropoff_datetime", "dropoff_datetime"),
        ("Airport_fee (Yellow only)", "airport_fee; NULL for Green"),
        ("trip_type (Green only)", "trip_type; NULL for Yellow"),
        ("ehail_fee", "Dropped because it carries no analytical value in observed data"),
    ])

    doc.add_heading("4.3 Data-Quality Validation", level=2)
    add_bullets(doc, [
        "Pickup and drop-off timestamps must be present.",
        "Drop-off cannot occur before pickup.",
        "Pickup and drop-off location IDs must be present.",
        "Trip distance must be present and non-negative.",
        "Fare amount and total amount must be present.",
        "Passenger count cannot be negative when supplied.",
        "Required source columns must exist before a batch proceeds.",
    ])
    doc.add_paragraph(
        "Invalid rows are written to pipeline.rejected_records with the trip key, combined "
        "rejection reason, JSON representation and run identifier. This quarantine approach "
        "preserves evidence while allowing valid rows in the same source file to proceed."
    )

    doc.add_heading("4.4 Chunked Processing", level=2)
    doc.add_paragraph(
        "PyArrow reads Parquet record batches using the configured CHUNK_SIZE (50,000 rows in "
        "the Compose environment). Only one chunk is converted to Pandas at a time, reducing "
        "Python memory requirements for Yellow files containing several million records. "
        "Accepted and rejected rows are inserted in pages of 2,000 database values."
    )

    doc.add_heading("4.5 Deduplication and Idempotency", level=2)
    doc.add_paragraph(
        "The source does not provide a universal trip identifier. The pipeline derives an MD5 "
        "trip_key from taxi type, pickup/drop-off timestamps, pickup/drop-off locations, distance, "
        "fare, total and vendor. Gold declares trip_key as its primary key. Promotion uses "
        "ON CONFLICT (trip_key) DO NOTHING, so an exact rerun does not create duplicate Gold rows."
    )
    add_callout(doc, "Important qualification", "The derived key is an inferred business key, not a source-issued identifier. Two legitimate trips with identical selected attributes could be treated as duplicates. MD5 is used for compact deterministic identity, not security.", AMBER)

    doc.add_heading("4.6 Transaction and Failure Model", level=2)
    doc.add_paragraph(
        "A run header is committed first so a failed attempt remains visible. Silver replacement, "
        "rejected rows, Gold promotion, materialized-view refresh and successful counters then "
        "execute in one monthly transaction. Any failure rolls that transaction back and preserves "
        "the last committed state. Database connection establishment waits through temporary "
        "unavailability, and the CLI retries interruptions."
    )
    add_callout(doc, "Tradeoff", "Atomic monthly publication simplifies correctness, but a restart near completion loses the full month of in-progress work and generates substantial WAL and disk I/O.", AMBER)

    doc.add_heading("4.7 Ingestion and Backfill", level=2)
    doc.add_paragraph(
        "The ingest command processes one month. Backfill expands an inclusive date range into "
        "monthly units and executes them sequentially. Backfill is historical reprocessing, not a "
        "database backup. Each month is independently audited and can be safely rerun."
    )
    add_code(doc, """
docker compose run --rm pipeline ingest --taxi-type green --year 2025 --month 1

docker compose run --rm pipeline backfill \
  --taxi-type green --start 2025-02 --end 2025-12
""")

    doc.add_heading("5. Database Design", level=1)
    add_table(doc, ["Schema", "Object", "Purpose"], [
        ("pipeline", "batch_runs", "One record per execution attempt, status, counts and errors"),
        ("pipeline", "file_ingestions", "Bronze path, SHA-256 checksum, status and row count"),
        ("pipeline", "rejected_records", "Invalid source rows and rejection reasons"),
        ("silver", "trips", "Canonical accepted records from both taxi types"),
        ("gold", "yellow_trips", "Deduplicated Yellow Taxi analytical records"),
        ("gold", "green_trips", "Deduplicated Green Taxi analytical records"),
        ("gold", "monthly_summary", "Materialized monthly counts, average distance and revenue"),
    ])
    doc.add_heading("5.1 Operational Metadata", level=2)
    doc.add_paragraph(
        "batch_runs uses running, success and failed statuses and records extracted, loaded and "
        "rejected counts. file_ingestions links the audit attempt to the exact Bronze path and "
        "checksum. Running counters currently remain zero until the monthly transaction commits; "
        "chunk progress is visible in logs rather than durable metadata."
    )
    doc.add_heading("5.2 Lineage", level=2)
    add_bullets(doc, [
        "File lineage: path, taxi type, source year/month and SHA-256 checksum.",
        "Row lineage: source filename, year/month and ingestion timestamp.",
        "Execution lineage: run ID, status, timestamps, counters and error message.",
        "Quality lineage: rejected row, reason, JSON record and originating run ID.",
    ])

    doc.add_heading("6. Configuration and Operation", level=1)
    doc.add_heading("6.1 Environment Variables", level=2)
    add_table(doc, ["Variable", "Purpose", "Compose value"], [
        ("POSTGRES_HOST", "Database hostname", "postgres"),
        ("POSTGRES_PORT", "Internal database port", "5432"),
        ("POSTGRES_DB", "Database name", "nyc_taxi"),
        ("POSTGRES_USER", "Database user", "taxi_user"),
        ("POSTGRES_PASSWORD", "Database password", "Environment secret"),
        ("BRONZE_DATA_PATH", "Raw storage root", "/app/data/bronze"),
        ("CHUNK_SIZE", "Rows per PyArrow batch", "50000"),
        ("LOG_LEVEL", "Runtime log verbosity", "INFO"),
    ])
    doc.add_heading("6.2 Start and Run", level=2)
    add_code(doc, """
cd "C:\\Users\\ekica\\Documents\\python pipeline\\Python-batch-pipeline"
docker compose up -d postgres
docker compose build pipeline
docker compose run --rm pipeline ingest --taxi-type green --year 2025 --month 1
""")
    doc.add_heading("6.3 pgAdmin Connection", level=2)
    add_table(doc, ["Setting", "Value"], [
        ("Host", "127.0.0.1"),
        ("Port", "5433"),
        ("Maintenance database", "nyc_taxi"),
        ("Username", "taxi_user"),
    ])
    doc.add_paragraph(
        "Port 5433 avoids the native Windows PostgreSQL server on port 5432. Materialized views "
        "appear under pgAdmin's Materialized Views node, not under Tables."
    )

    doc.add_heading("7. Testing and Verification", level=1)
    doc.add_heading("7.1 Automated Tests", level=2)
    doc.add_paragraph("The unit suite produced seven passing tests at the recorded verification checkpoint.")
    add_bullets(doc, [
        "Download URL and Bronze path construction",
        "Month ranges crossing a year boundary",
        "Invalid month handling",
        "Canonical Yellow transformation",
        "Deterministic trip-key generation",
        "Accepted/rejected row splitting",
        "Required-column detection",
    ])
    add_code(doc, "python -m pytest tests -q -p no:cacheprovider\n\nExpected: 7 passed")
    doc.add_heading("7.2 Idempotency Test", level=2)
    add_table(doc, ["Execution", "Extracted", "Loaded", "New Gold rows"], [
        ("Initial Green January run", "48,326", "48,326", "48,326"),
        ("Immediate rerun", "48,326", "48,326", "0"),
    ])
    doc.add_paragraph(
        "Gold remained at 48,326 Green January rows after both attempts. This verifies that the "
        "same Bronze file can be reprocessed while the Gold primary key prevents duplication."
    )
    doc.add_heading("7.3 Reconciliation", level=2)
    doc.add_paragraph(
        "For every successful latest monthly run, extracted_rows equals loaded_rows plus "
        "rejected_rows. Showcase SQL also verifies checksum format, lineage completeness, Gold "
        "primary-key constraints and Silver-to-Gold counts."
    )

    doc.add_heading("8. Results", level=1)
    add_callout(doc, "Snapshot basis", "Docker Desktop was not running when this Word report was generated. Results below use the latest verified database queries and evidence recorded on 31 July 2026; they are explicitly presented as a checkpoint.", LIGHT_BLUE)
    add_table(doc, ["Taxi type", "Completed coverage", "Accepted rows", "Rejected rows"], [
        ("Green", "January-December 2025", "590,031", "1,344"),
        ("Yellow", "January-March 2025", "11,197,728", "298"),
        ("Combined", "15 taxi-month batches", "11,787,759", "1,642"),
    ])
    doc.add_paragraph(
        "The row totals above use the latest successful run per taxi type and month. Operational "
        "history contains additional successful rerun attempts, but they must not be summed as new "
        "business data. Gold remains deduplicated."
    )
    doc.add_heading("8.1 Observed Data Quality", level=2)
    doc.add_paragraph(
        "The most frequently observed rejection was 'dropoff before pickup'. Yellow batches had "
        "very low rejection rates, while several early Green months rejected roughly 0.45%-0.60% "
        "of source rows. Rejected records remain queryable for root-cause analysis."
    )
    doc.add_heading("8.2 Reporting Output", level=2)
    doc.add_paragraph(
        "The populated gold.monthly_summary materialized view contains one row for each loaded "
        "taxi type and month, with trip count, average distance and total revenue. At the verified "
        "checkpoint it contained 15 rows: twelve Green months and three Yellow months."
    )

    doc.add_heading("9. Incidents and Lessons Learned", level=1)
    doc.add_heading("9.1 PostgreSQL Restart During Yellow February", level=2)
    doc.add_paragraph(
        "PostgreSQL received an administrator-triggered fast shutdown while a monthly transaction "
        "was active. The connection closed, the transaction rolled back and subsequent attempts "
        "initially encountered DNS, connection-refused and recovery messages. The incident led to "
        "connection waiting, interrupted-batch retries and explicit recovery documentation."
    )
    doc.add_heading("9.2 Quiet Logs During Yellow March", level=2)
    doc.add_paragraph(
        "After all 4,145,257 Yellow March source rows had been processed, Python emitted no new "
        "chunk logs while PostgreSQL promoted 4,145,176 valid rows into Gold. Database activity "
        "showed WALWrite and DataFileWrite waits. The key lesson is to inspect pg_stat_activity and "
        "database logs before declaring a bulk operation stalled."
    )
    doc.add_heading("9.3 Local PostgreSQL Port Conflict", level=2)
    doc.add_paragraph(
        "pgAdmin initially connected to a different native PostgreSQL service on host port 5432, "
        "where the nyc_taxi database did not exist. Publishing Docker PostgreSQL on port 5433 "
        "resolved the ambiguity without changing the container's internal port."
    )

    doc.add_heading("10. Limitations and Critical Evaluation", level=1)
    limitations = [
        ("Large monthly transaction", "A restart near completion rolls back the full month.", "Use durable run-scoped staging and a smaller publish transaction."),
        ("Insert-heavy Silver/Gold flow", "Rows and indexes are written twice and create substantial WAL.", "Use COPY, partitioning and bulk-load tuning."),
        ("No resumable checkpoints", "A failed month is reread from the beginning.", "Persist Parquet row-group or chunk checkpoints."),
        ("Full summary refresh", "Reporting refresh cost grows with Gold.", "Incrementally update affected months or schedule refresh separately."),
        ("Approximate trip key", "Identical legitimate trips may be merged.", "Agree source identity semantics or adopt a future source ID."),
        ("Schema evolution is manual", "New useful source fields are dropped until code and SQL change.", "Add schema-version contracts, drift alerts and migrations."),
        ("Validation is incomplete", "Zone validity, plausible speed/duration and financial bounds are unchecked.", "Create versioned business-quality rules and metrics."),
        ("No concurrency lock", "Two workers can process the same month simultaneously.", "Use PostgreSQL advisory locks or active-run uniqueness."),
        ("Initialization, not migrations", "SQL edits do not update an existing Docker volume.", "Adopt Alembic, Flyway or an equivalent migration tool."),
        ("Limited observability", "No durable heartbeat, metrics or alerting.", "Add structured logs, metrics and stale-run alerts."),
        ("Limited automated integration testing", "Rollback and database promotion are not covered in CI.", "Add disposable PostgreSQL integration and failure-injection tests."),
        ("Development-grade security", "Local credentials, exposed port and no TLS or role separation.", "Use secret management, least privilege, private networking and TLS."),
        ("Local Bronze storage", "Raw data has no object-store durability or lifecycle management.", "Move Bronze to S3, Azure Blob or GCS."),
    ]
    add_table(doc, ["Limitation", "Impact", "Recommended improvement"], limitations, [1.6, 2.4, 2.8])

    doc.add_heading("11. Recommended Roadmap", level=1)
    add_numbered(doc, [
        "Replace execute_values with PostgreSQL COPY into durable run-scoped staging tables.",
        "Add advisory locking and a unique policy for active taxi-type/month batches.",
        "Introduce formal database migrations and automated disposable-PostgreSQL integration tests.",
        "Add HTTP retry/backoff, Parquet integrity checks and schema-drift alerts.",
        "Persist chunk heartbeats and checkpoints so interrupted batches can resume.",
        "Replace full materialized-view refreshes with incremental monthly aggregation.",
        "Add a scheduler, structured observability, alerts and production secret management.",
        "Move Bronze files to durable cloud object storage for non-local deployment.",
    ])

    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "The NYC Taxi Batch Pipeline meets its core objective: it provides a functioning, "
        "auditable and repeatable monthly Python ingestion workflow across heterogeneous Yellow "
        "and Green Taxi source schemas. The project demonstrates practical batch engineering "
        "through raw-data preservation, bounded-memory processing, canonical transformation, "
        "quality quarantine, transactional loading, metadata, lineage, backfill support and "
        "idempotent Gold publication."
    )
    doc.add_paragraph(
        "Its strongest engineering qualities are transparency and recoverability. Source files, "
        "checksums, run attempts, rejected rows and final records are traceable. Its principal "
        "weakness is scalability of the monthly transaction and insert path. The proposed roadmap "
        "provides a credible progression from a successful local portfolio project to a more "
        "resumable, observable and production-oriented data platform."
    )

    doc.add_page_break()
    doc.add_heading("Appendix A: Repository Structure", level=1)
    add_code(doc, """
Python-batch-pipeline/
|-- taxi_pipeline/
|   |-- __main__.py       CLI and backfill control
|   |-- config.py         Environment configuration
|   |-- extract.py        Source download and Bronze paths
|   |-- transform.py      Canonical schema and validation
|   |-- database.py       Transactions and PostgreSQL loading
|   `-- pipeline.py       End-to-end monthly orchestration
|-- sql/                  Schema and table initialization
|-- tests/                Unit tests
|-- showcase/             Read-only demonstration queries
|-- docs/                 Detailed project documentation
|-- data/bronze/          Raw monthly Parquet files
|-- Dockerfile
|-- docker-compose.yml
`-- requirements.txt
""")

    doc.add_heading("Appendix B: Useful Operational SQL", level=1)
    add_code(doc, """
-- Recent runs
SELECT run_id, taxi_type, source_year, source_month, status,
       extracted_rows, loaded_rows, rejected_rows
FROM pipeline.batch_runs
ORDER BY run_id DESC;

-- Gold counts
SELECT 'yellow' AS taxi_type, COUNT(*) FROM gold.yellow_trips
UNION ALL
SELECT 'green', COUNT(*) FROM gold.green_trips;

-- Monthly report
SELECT * FROM gold.monthly_summary
ORDER BY source_year, source_month, taxi_type;

-- Database activity during a quiet bulk operation
SELECT state, wait_event_type, wait_event,
       NOW() - query_start AS age, LEFT(query, 120) AS query
FROM pg_stat_activity
WHERE datname = 'nyc_taxi' AND state <> 'idle';
""")

    doc.add_heading("Appendix C: Idempotency Demonstration", level=1)
    add_code(doc, """
$before = [int64](docker exec nyc_taxi_postgres psql \
  -U taxi_user -d nyc_taxi -At \
  -c "SELECT COUNT(*) FROM gold.green_trips
      WHERE source_year=2025 AND source_month=1;")

docker compose run --rm pipeline ingest \
  --taxi-type green --year 2025 --month 1

$after = [int64](docker exec nyc_taxi_postgres psql \
  -U taxi_user -d nyc_taxi -At \
  -c "SELECT COUNT(*) FROM gold.green_trips
      WHERE source_year=2025 AND source_month=1;")

if ($before -ne $after) { throw "Idempotency test failed" }
Write-Host "Idempotency test passed" -ForegroundColor Green
""")

    doc.add_heading("Appendix D: Supporting Documentation", level=1)
    add_bullets(doc, [
        "README.md - setup and quick-start guide",
        "docs/architecture.md - detailed architecture and transaction behavior",
        "docs/data-model.md - schemas, tables, keys and lineage",
        "docs/operations.md - monitoring and recovery runbook",
        "docs/limitations.md - critical evaluation and roadmap",
        "docs/decision-log.md - architecture decisions and tradeoffs",
        "docs/evidence.md - observed test and runtime evidence",
        "showcase/ - validated SQL presentation queries",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
